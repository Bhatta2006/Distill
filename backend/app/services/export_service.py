"""
Export engine: flatten highlights + drawings onto PDF, append summary page + diagram.
Also handles Markdown and DOCX formats.
"""
from __future__ import annotations
from pathlib import Path
import fitz  # PyMuPDF

from app.models.document import Document, Section
from app.models.annotation import Highlight, Drawing
from app.models.diagram import Diagram


def export_pdf(
    original_path: str,
    highlights: list[Highlight],
    drawings: list[Drawing],
    sections: list[Section],
    diagram: Diagram | None,
    output_path: str,
) -> None:
    doc = fitz.open(original_path)

    # Layer 1: AI + user highlights
    _apply_highlights(doc, highlights)

    # Layer 2: Freehand drawings
    _apply_drawings(doc, drawings)

    # Append summary page
    _append_summary_page(doc, sections)

    # Append diagram page
    if diagram:
        _append_diagram_page(doc, diagram)

    doc.save(output_path)
    doc.close()


def _apply_highlights(doc: fitz.Document, highlights: list[Highlight]) -> None:
    color_map = {
        "#FFD700": (1.0, 0.84, 0.0),
        "#90EE90": (0.56, 0.93, 0.56),
        "#ADD8E6": (0.68, 0.85, 0.90),
        "#FFB6C1": (1.0, 0.71, 0.76),
    }

    for h in highlights:
        page_idx = h.page - 1
        if page_idx < 0 or page_idx >= len(doc):
            continue
        page = doc[page_idx]
        rgb = color_map.get(h.color, (1.0, 0.84, 0.0))
        for rect_data in h.rects:
            rect = fitz.Rect(
                rect_data["x"],
                rect_data["y"],
                rect_data["x"] + rect_data["width"],
                rect_data["y"] + rect_data["height"],
            )
            annot = page.add_highlight_annot(rect)
            annot.set_colors(stroke=rgb)
            annot.update()
            if h.comment:
                page.add_text_annot(rect.tl, h.comment)


def _apply_drawings(doc: fitz.Document, drawings: list[Drawing]) -> None:
    for drawing in drawings:
        page_idx = drawing.page - 1
        if page_idx < 0 or page_idx >= len(doc):
            continue
        # ponytail: Fabric.js paths are complex; we just note the drawing exists via a stamp
        # Full SVG rendering would require a headless browser — deferred to Phase 2
        page = doc[page_idx]
        page.add_text_annot(fitz.Point(10, 10), "[Freehand drawing — view in Lucid]")


def _append_summary_page(doc: fitz.Document, sections: list[Section]) -> None:
    page = doc.new_page(width=612, height=792)
    tf = page.insert_textbox(
        fitz.Rect(40, 40, 572, 752),
        "LUCID — PAPER SUMMARY\n\n" + "\n\n".join(
            f"{s.title.upper()}\n{s.summary or 'No summary available.'}"
            for s in sections if s.summary
        ),
        fontsize=10,
        fontname="helv",
        color=(0, 0, 0),
    )


def _append_diagram_page(doc: fitz.Document, diagram: Diagram) -> None:
    page = doc.new_page(width=612, height=792)
    page.insert_textbox(
        fitz.Rect(40, 40, 572, 752),
        f"LUCID — METHOD DIAGRAM (Confidence: {diagram.confidence.upper()})\n\n"
        f"[Mermaid.js source — paste at mermaid.live to render]\n\n{diagram.mermaid_code}",
        fontsize=9,
        fontname="cour",
        color=(0, 0, 0),
    )


def export_markdown(sections: list[Section], diagram: Diagram | None) -> str:
    parts = ["# Paper Summary — Lucid\n"]
    for s in sections:
        parts.append(f"## {s.title}\n{s.summary or '_No summary_'}\n")
    if diagram:
        parts.append(f"## Method Diagram\n\n```mermaid\n{diagram.mermaid_code}\n```\n")
    return "\n".join(parts)


def export_docx(sections: list[Section], diagram: Diagram | None, output_path: str) -> None:
    from docx import Document as DocxDoc
    doc = DocxDoc()
    doc.add_heading("Paper Summary — Lucid", 0)
    for s in sections:
        doc.add_heading(s.title, level=2)
        doc.add_paragraph(s.summary or "No summary available.")
    if diagram:
        doc.add_heading("Method Diagram (Mermaid source)", level=2)
        doc.add_paragraph(diagram.mermaid_code)
    doc.save(output_path)
